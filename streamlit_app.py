import os
import io
import uuid
import requests
import streamlit as st

# ------------------------------------------------------------------
# Load secrets (API keys) from Streamlit Cloud's secrets manager
# ------------------------------------------------------------------
for key in ["OPENAI_API_KEY", "GROQ_API_KEY", "QDRANT_URL", "QDRANT_API_KEY", "LLM_PROVIDER", "API_KEY_HERE"]:
    if key in st.secrets:
        os.environ[key] = st.secrets[key]

from main import orchestrator

# ------------------------------------------------------------------
# Page setup + custom styling
# ------------------------------------------------------------------
st.set_page_config(page_title="AI Resume & Cover Letter Generator", page_icon="📄", layout="centered")

st.markdown("""
<style>
.hero {
    background: linear-gradient(90deg, #4F46E5 0%, #9333EA 100%);
    padding: 2rem 1.5rem 1.5rem 1.5rem;
    border-radius: 16px;
    color: white;
    margin-bottom: 0.75rem;
}
.hero h1 { margin: 0; font-size: 2.2rem; }
.hero p { margin-top: 0.4rem; opacity: 0.9; }
.badge-row { margin-bottom: 1.2rem; }
.badge {
    display: inline-block;
    background: #F3F0FF;
    color: #4F46E5;
    border-radius: 999px;
    padding: 0.25rem 0.8rem;
    font-size: 0.8rem;
    font-weight: 600;
    margin-right: 0.4rem;
    margin-bottom: 0.4rem;
}
.privacy-note {
    font-size: 0.85rem;
    color: #6B7280;
    text-align: center;
    margin-top: 0.5rem;
}
div.stButton > button, div.stDownloadButton > button {
    background: linear-gradient(90deg, #4F46E5 0%, #9333EA 100%);
    color: white;
    border: none;
    border-radius: 8px;
    padding: 0.6rem 1.4rem;
    font-weight: 600;
}
div.stButton > button:hover, div.stDownloadButton > button:hover {
    opacity: 0.9;
    color: white;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="hero">
    <h1>📄 AI Resume & Cover Letter Generator</h1>
    <p>Upload your resume and a job description — a multi-agent AI pipeline tailors your resume, scores it for ATS match, and drafts a matching cover letter, all in one pass.</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<style>
.coffee-loader {
    text-align: center;
    padding: 1rem 0 0.5rem 0;
}
.coffee-loader .cup {
    font-size: 3rem;
    display: inline-block;
    animation: coffee-bounce 1.4s ease-in-out infinite;
}
.coffee-loader p {
    color: #6B7280;
    font-weight: 600;
    margin-top: 0.3rem;
}
@keyframes coffee-bounce {
    0%, 100% { transform: translateY(0) rotate(0deg); }
    50% { transform: translateY(-12px) rotate(-6deg); }
}
</style>
""", unsafe_allow_html=True)

st.markdown(
    '<div class="badge-row">'
    '<span class="badge">⚡ Groq LLM</span>'
    '<span class="badge">🧠 Multi-Agent</span>'
    '<span class="badge">🔍 RAG Matching</span>'
    '<span class="badge">🐍 FastAPI</span>'
    '</div>',
    unsafe_allow_html=True,
)

# ------------------------------------------------------------------
# Usage counter (free, external, persists across app restarts)
# ------------------------------------------------------------------
COUNTER_NAMESPACE = "ai-resume-generator-demo"
COUNTER_KEY = "resumes-generated"

def get_counter_count():
    try:
        resp = requests.get(f"https://api.counterapi.dev/v1/{COUNTER_NAMESPACE}/{COUNTER_KEY}", timeout=3)
        return resp.json().get("count", None)
    except Exception:
        return None

def increment_counter():
    try:
        resp = requests.get(f"https://api.counterapi.dev/v1/{COUNTER_NAMESPACE}/{COUNTER_KEY}/up", timeout=3)
        return resp.json().get("count", None)
    except Exception:
        return None

current_count = get_counter_count()
if current_count is not None:
    st.caption(f"✨ {current_count} resumes generated so far by people using this tool")

# ------------------------------------------------------------------
# How it works
# ------------------------------------------------------------------
with st.expander("ℹ️ How this works (5-agent pipeline)"):
    st.markdown("""
1. **Profile Analyzer** — reads your background and experience level
2. **ATS Optimizer** — scores your resume against the job description using semantic + keyword matching
3. **Resume Writer** — drafts a tailored first version
4. **Reviewer** — checks grammar, formatting, and consistency
5. **Human Optimizer** — polishes the final version to sound natural, not AI-generated
6. **Cover Letter Writer** — drafts a matching cover letter using the same context

All steps run automatically in sequence — you'll see live progress below once you click Generate.
    """)

# ------------------------------------------------------------------
# File parsing helpers
# ------------------------------------------------------------------
def extract_text_from_upload(uploaded_file):
    if uploaded_file is None:
        return None
    name = uploaded_file.name.lower()

    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(uploaded_file)
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    if name.endswith(".docx"):
        import docx
        document = docx.Document(uploaded_file)
        return "\n".join(p.text for p in document.paragraphs)

    return None


def build_docx_bytes(text: str) -> bytes:
    import docx
    document = docx.Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(text: str) -> bytes:
    from fpdf import FPDF

    # Core PDF fonts only support latin-1 — sanitize common unicode punctuation
    replacements = {
        "\u2014": "-", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2022": "-", "\u2026": "...",
    }
    safe_text = text
    for old, new in replacements.items():
        safe_text = safe_text.replace(old, new)
    safe_text = safe_text.encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    for line in safe_text.split("\n"):
        pdf.multi_cell(0, 6, line)
    return bytes(pdf.output())


SAMPLE_RESUME = """Jordan Lee
Software Engineer
jordan.lee@email.com | (555) 123-4567 | Boston, MA

PROFESSIONAL SUMMARY
Backend Software Engineer with 5 years of experience building REST APIs and cloud-native
applications. Strong background in Python, FastAPI, and AWS.

EXPERIENCE
TechCorp Inc. — Remote | Jan 2021 - Present
Software Engineer
- Built and maintained REST APIs serving 1M+ daily requests using Python and FastAPI
- Deployed microservices on AWS Lambda and ECS, reducing infra costs by 20%
- Automated CI/CD pipelines using GitHub Actions and Docker

SKILLS
Python, FastAPI, AWS, Docker, PostgreSQL, REST APIs, Git
"""

SAMPLE_JD = """We are hiring a Backend Software Engineer with experience in Python, FastAPI, and AWS.
The ideal candidate has built and scaled REST APIs, has experience with Docker and CI/CD,
and is comfortable working in an Agile team environment.
"""

# ------------------------------------------------------------------
# Sample data button
# ------------------------------------------------------------------
if "full_name" not in st.session_state:
    st.session_state.full_name = ""
    st.session_state.current_role = ""
    st.session_state.skills_input = ""
    st.session_state.experience_years = 0
    st.session_state.resume_text_area = ""
    st.session_state.job_description_area = ""

if st.button("🎬 Try with sample data instead"):
    st.session_state.full_name = "Jordan Lee"
    st.session_state.current_role = "Software Engineer"
    st.session_state.skills_input = "Python, FastAPI, AWS, Docker, PostgreSQL, REST APIs, Git"
    st.session_state.experience_years = 5
    st.session_state.resume_text_area = SAMPLE_RESUME
    st.session_state.job_description_area = SAMPLE_JD
    st.rerun()

# ------------------------------------------------------------------
# Input section
# ------------------------------------------------------------------
with st.container(border=True):
    st.subheader("1. Your Details")
    col1, col2 = st.columns(2)
    with col1:
        full_name = st.text_input("Full Name", key="full_name")
        experience_years = st.number_input("Years of Experience", min_value=0, max_value=50, step=1, key="experience_years")
    with col2:
        current_role = st.text_input("Current Role", placeholder="e.g. Software Engineer", key="current_role")
        skills_input = st.text_input("Skills (comma-separated)", placeholder="Python, FastAPI, AWS", key="skills_input")

with st.container(border=True):
    st.subheader("2. Your Resume")
    input_mode = st.radio("How would you like to provide your resume?", ["Upload a file", "Paste text"], horizontal=True)

    resume_text = None
    if input_mode == "Upload a file":
        uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
        if uploaded_file is not None:
            resume_text = extract_text_from_upload(uploaded_file)
            if resume_text:
                st.success(f"Loaded {uploaded_file.name} ({len(resume_text)} characters)")
            else:
                st.error("Couldn't read that file — try TXT or DOCX instead.")
        elif st.session_state.resume_text_area:
            resume_text = st.session_state.resume_text_area
            st.info("Using sample resume text (switch to 'Paste text' to view/edit it).")
    else:
        resume_text = st.text_area("Paste your resume text here", height=220, key="resume_text_area")

with st.container(border=True):
    st.subheader("3. Job Description (optional, enables smart matching + tailored cover letter)")
    job_description = st.text_area("Paste the job description here", height=150, key="job_description_area")

st.write("")

MAX_GENERATIONS_PER_SESSION = 3
if "generation_count" not in st.session_state:
    st.session_state.generation_count = 0

remaining = MAX_GENERATIONS_PER_SESSION - st.session_state.generation_count
if remaining <= 0:
    st.warning(
        f"You've reached the limit of {MAX_GENERATIONS_PER_SESSION} generations for this session "
        "(this keeps the free demo available for everyone). Refresh the page to reset, or try again later."
    )
    generate_clicked = False
else:
    st.caption(f"{remaining} generation(s) remaining this session")
    generate_clicked = st.button("✨ Generate Resume + Cover Letter", use_container_width=True)

st.markdown('<p class="privacy-note">🔒 Your resume text isn\'t stored anywhere — it only exists for the duration of this session.</p>', unsafe_allow_html=True)

# ------------------------------------------------------------------
# Run pipeline with live progress
# ------------------------------------------------------------------
if generate_clicked:
    if not full_name or not resume_text:
        st.error("Please provide at least your name and a resume (uploaded or pasted).")
    else:
        coffee_placeholder = st.empty()
        coffee_placeholder.markdown("""
        <div class="coffee-loader">
            <div class="cup">☕</div>
            <p>Brewing your resume... grab a coffee, this takes about 20-30 seconds</p>
        </div>
        """, unsafe_allow_html=True)

        with st.status("Starting AI agents...", expanded=True) as status:
            def update(msg):
                status.update(label=msg)
                st.write(f"→ {msg}")

            user_request = {
                "api_key": "internal",
                "full_name": full_name,
                "current_role": current_role,
                "skills": [s.strip() for s in skills_input.split(",") if s.strip()],
                "experience_years": int(experience_years),
                "resume_text": resume_text,
                "resume_file": None,
                "job_description": job_description or None,
            }
            request_id = str(uuid.uuid4())
            result = orchestrator(user_request, request_id, progress_callback=update)
            status.update(label="All agents finished!", state="complete")

        coffee_placeholder.empty()

        st.success(f"Done! (took {result['execution_time']}s)")
        increment_counter()
        st.session_state.generation_count += 1

        # Store the final resume in session state so Apply-fix buttons can edit it
        # and have the edits persist across reruns.
        st.session_state.editable_resume = result["workflow"]["human_optimizer"]["human_friendly_resume"]
        st.session_state.cover_letter = result["workflow"]["cover_letter"]["cover_letter"]
        st.session_state.last_result = result
        st.session_state.last_full_name = full_name

# ------------------------------------------------------------------
# Display results (persists across reruns via session_state, so Apply buttons work)
# ------------------------------------------------------------------
if "last_result" in st.session_state:
    result = st.session_state.last_result
    full_name = st.session_state.last_full_name
    cover_letter = st.session_state.cover_letter

    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        ["✅ Final Resume", "✉️ Cover Letter", "📝 First Draft", "📊 ATS Score", "🔍 Reviewer Notes"]
    )

    with tab1:
        completeness = result["workflow"].get("completeness_check", {})
        if completeness.get("possibly_missing"):
            with st.container(border=True):
                st.warning(
                    f"⚠️ Content check: {completeness['completeness_pct']}% of detected resume entries "
                    "appear to be present. The following original line(s) weren't found in the final "
                    "version — please double-check nothing important was dropped:"
                )
                for line in completeness["possibly_missing"]:
                    st.caption(f"• {line}")

        st.write(st.session_state.editable_resume)
        dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)
        with dl_col1:
            st.download_button(
                "⬇️ TXT",
                data=st.session_state.editable_resume,
                file_name=f"{full_name.replace(' ', '_')}_resume.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with dl_col2:
            st.download_button(
                "⬇️ DOCX",
                data=build_docx_bytes(st.session_state.editable_resume),
                file_name=f"{full_name.replace(' ', '_')}_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with dl_col3:
            st.download_button(
                "⬇️ PDF",
                data=build_pdf_bytes(st.session_state.editable_resume),
                file_name=f"{full_name.replace(' ', '_')}_resume.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        with dl_col4:
            if st.button("↩️ Reset edits", use_container_width=True):
                st.session_state.editable_resume = result["workflow"]["human_optimizer"]["human_friendly_resume"]
                st.rerun()

    with tab2:
        st.write(cover_letter)
        dl_col1, dl_col2, dl_col3 = st.columns(3)
        with dl_col1:
            st.download_button(
                "⬇️ TXT",
                data=cover_letter,
                file_name=f"{full_name.replace(' ', '_')}_cover_letter.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with dl_col2:
            st.download_button(
                "⬇️ DOCX",
                data=build_docx_bytes(cover_letter),
                file_name=f"{full_name.replace(' ', '_')}_cover_letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with dl_col3:
            st.download_button(
                "⬇️ PDF",
                data=build_pdf_bytes(cover_letter),
                file_name=f"{full_name.replace(' ', '_')}_cover_letter.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

    with tab3:
        st.write(result["workflow"]["resume_writer"]["generated_resume"])

    with tab4:
        ats_data = result["workflow"]["ats_optimization"]
        score = ats_data.get("ats_score", 0)
        st.metric("ATS Match Score", f"{score}/100")
        st.progress(score / 100)
        with st.expander("Detailed feedback"):
            st.write(ats_data.get("llm_feedback", ""))

    with tab5:
        suggestions = result["workflow"]["reviewer"].get("suggestions", [])
        if not suggestions:
            st.write(result["workflow"]["reviewer"].get("review_feedback", "No issues found."))
        else:
            st.caption("Click Apply to instantly update the Final Resume tab with a fix.")
            for i, sug in enumerate(suggestions):
                issue = sug.get("issue", "Suggestion")
                current_text = sug.get("current_text", "")
                suggested_fix = sug.get("suggested_fix", "")

                with st.container(border=True):
                    st.markdown(f"**{issue}**")
                    col_a, col_b, col_c = st.columns([2, 2, 1])
                    with col_a:
                        st.caption("Current")
                        st.code(current_text, language=None)
                    with col_b:
                        st.caption("Suggested")
                        st.code(suggested_fix, language=None)
                    with col_c:
                        st.write("")
                        already_applied = current_text and current_text not in st.session_state.editable_resume
                        if already_applied:
                            st.success("Applied")
                        elif st.button("Apply", key=f"apply_{i}", use_container_width=True):
                            if current_text and current_text in st.session_state.editable_resume:
                                st.session_state.editable_resume = st.session_state.editable_resume.replace(
                                    current_text, suggested_fix, 1
                                )
                                st.rerun()
                            else:
                                st.warning("Couldn't find an exact match to auto-apply — edit manually.")

# ------------------------------------------------------------------
# Feedback footer
# ------------------------------------------------------------------
st.divider()
st.markdown(
    "### Found this useful? \n"
    "[⭐ Leave a quick review here](PASTE_YOUR_GOOGLE_FORM_LINK_HERE) — it takes 30 seconds and helps a lot!"
)
